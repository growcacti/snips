import tkinter as tk
from tkinter import ttk
from tkinter import ttk
from tkinter import filedialog
import keyboard
class Themes:
    def __init__(self):
        self.Theme = tk.Menu(menubar.toolbar, tearoff=False)
        self.themes = tk.Menu(self.Theme, tearoff=False)

    def theme_menu(self):
        self.Theme.add_command(label=" Text Color", command=lambda: self.text_color())
        self.Theme.add_command(
            label=" Cursor Color", command=lambda: self.cursor_color()
        )
        self.Theme.add_command(
            label=' Selector" Color', command=lambda: self.selector_color()
        )
        self.themes.add_radiobutton(
            label="White (Default)", command=lambda: self.white_theme()
        )
        self.themes.add_radiobutton(label="White", command=lambda: self.white_theme())
        self.themes.add_radiobutton(
            label="Valhalla", command=lambda: self.valhalla_theme()
        )
        self.themes.add_radiobutton(label="Havana", command=lambda: self.havana_theme())
        menubar.toolbar.add_cascade(label="Theme", menu=self.Theme)
        self.Theme.add_cascade(label="  Themes", menu=self.themes)
        w.root.config(menu=self.Theme)

    def dark_theme(self):
        s = ttk.Style(w.root)
        s.configure(style="TNotebook", background="#282828")
        w.root.configure(bg="#282828")
        txtconfig.background = "#282828"
        txtconfig.foreground = "white"
        txtconfig.insertbackground = "white"

        for i in w.tab.txt_collection:
            i.config(
                background=txtconfig.background,
                foreground=txtconfig.foreground,
                insertbackground=txtconfig.insertbackground,
            )
            w.root.update()

    def valhalla_theme(self):
        s = ttk.Style(w.root)
        s.configure(style="TNotebook", background="#262442")
        w.root.configure(bg="#262442")
        txtconfig.background = "#262442"
        txtconfig.foreground = "white"
        txtconfig.insertclass Window:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("DK Don't Know")
        self.label = tk.Label(
            self.root,
            width=self.root.winfo_reqwidth(),
            height=self.root.winfo_reqheight(),
            text="Create new file",
        )
        self.tab = Tab(self.root)
        self.txt, self.tb = self.tab.add_tab("untitled")
        self.height, self.width = (
            self.root.winfo_reqheight(),
            self.root.winfo_reqwidth(),
        )
        self.txt.config(
            height=self.height,
            width=self.width,
        )
        self.footer = tk.Label(self.root)

        self.footer.pack(fill="both", expand="yes", side=tk.BOTTOM)
        self.txt.pack(fill="both", expand="yes")
        self.tb.pack(fill="both", expand="yes")
        self.label.pack()
        self.root.geometry("998x646")
        self.root.minsize(width=400, height=200)

    def footer_elements(self):
        while True:
            try:
                for i in w.tab.txt_collection:
                    word_count = len(i.get("1.0", tk.END).split())
                    character_count = len(i.get("1.0", tk.END)) - i.get(
                        "1.0", tk.END
                    ).count("\n")
                    mouse_pos = i.index(tk.INSERT)
                    line, column = (int(num) for num in mouse_pos.split("."))
                    self.footer.configure(
                        text=(
                            f"word count: {word_count}   "
                            f"character count: {character_count}     "
                            f"lines number: {line}    "
                            f"columns number: {column}"
                        ),
                        anchor=tk.E,
                    )
                self.root.update()
            except (tk.TclError, ValueError):
                break


w = Window()

background = "white"
        for i in w.tab.txt_collection:
            i.config(
                background=txtconfig.background,
                foreground=txtconfig.foreground,
                insertbackground=txtconfig.insertbackground,
            )
            w.root.update()

    def white_theme(self):
        s = ttk.Style(w.root)
        s.configure(style="TNotebook", background="white")
        w.root.configure(bg="white")
        txtconfig.background = "white"
        txtconfig.foreground = "black"
        txtconfig.insertbackground = "black"
        for i in w.tab.txt_collection:
            i.config(
                background=txtconfig.background,
                foreground=txtconfig.foreground,
                insertbackground=txtconfig.insertbackground,
            )
            w.root.update()

    def havana_theme(self):
        s = ttk.Style(w.root)
        w.root.configure(bg="#412A2B")
        s.configure("TNotebook", background="#412A2B")
        txtconfig.background = "#412A2B"
        txtconfig.foreground = "white"
        txtconfig.insertbackground = "white"
        for i in w.tab.txt_collection:
            i.config(
                background=txtconfig.background,
                foreground=txtconfig.foreground,
                insertbackground=txtconfig.insertbackground,
            )
            w.root.update()

    def text_color(self):
        txtconfig.foreground = colorchooser.askcolor()[1]
        for i in w.tab.txt_collection:
            i.config(foreground=txtconfig.foreground)

    def cursor_color(self):
        txtconfig.insertbackground = colorchooser.askcolor()[1]
        for i in w.tab.txt_collection:
            i.config(insertbackground=txtconfig.insertbackground)

    def selector_color(self):
        txtconfig.selectbackground = colorchooser.askcolor()[1]
        for i in w.tab.txt_collection:
            i.config(selectbackground=txtconfig.selectbackground)


# control any background color through txtconfig object

class Window:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("DK Don't Know")
        self.label = tk.Label(
            self.root,
            width=self.root.winfo_reqwidth(),
            height=self.root.winfo_reqheight(),
            text="Create new file",
        )
        self.tab = Tab(self.root)
        self.txt, self.tb = self.tab.add_tab("untitled")
        self.height, self.width = (
            self.root.winfo_reqheight(),
            self.root.winfo_reqwidth(),
        )
        self.txt.config(
            height=self.height,
            width=self.width,
        )
        self.footer = tk.Label(self.root)

        self.footer.pack(fill="both", expand="yes", side=tk.BOTTOM)
        self.txt.pack(fill="both", expand="yes")
        self.tb.pack(fill="both", expand="yes")
        self.label

class File:
    def __init__(self):
        self.filemenu = tk.Menu(menubar.toolbar, tearoff=False)
        self.binding_keys()

    def file_menu(self):
        self.filemenu.add_command(
            label="New file", command=lambda: self.new_file(), accelerator="Ctrl+N"
        )
        self.filemenu.add_command(
            label="Open File", command=lambda: self.open_file(), accelerator="Ctrl+O"
        )
        self.filemenu.add_command(
            label="Save", command=lambda: self.save_file(), accelerator="Ctrl+S"
        )
        self.filemenu.add_command(
            label="Save as",
            command=lambda: self.save_file_as(),
            accelerator="Ctrl+Shift+S",
        )
        self.filemenu.add_command(
            label="Close file", command=lambda: self.close_file(), accelerator="Ctrl+K"
        )
        self.filemenu.add_command(label="Exit", command=lambda: self.exit())
        menubar.toolbar.add_cascade(label="File", menu=self.filemenu)

    def open_file(self):
        try:
            self.f = filedialog.askopenfile(
                mode="r",
                initialdir="/",
                title="Select A File",
                filetypes=(("files", "*.txt"), ("all files", "*.*")),
            )
            filename = os.path.basename(self.f.name)
            txt, tb = w.tab.add_tab(f"{filename}")
            txt.config(
                height=w.height,
                width=w.width,
            )
            txt.insert("1.0", self.f.read())
            txt.pack()
            tb.pack()

        except AttributeError:
            return

    def save_file_as(self):
        try:
            for i in w.tab.txt_collection:
                self.f = filedialog.asksaveasfile(
                    mode="w",
                    defaultextension=".txt",
                    initialdir="/",
                    title="Select A File",
                    filetypes=(
                        ("text", "*.txt"),
                        ("python", "*.py"),
                        ("all files", "*.*"),
                    ),
                )
                self.f.write(i.get("1.0", tk.END))
                self.f.close()
                w.tab.notebook.tab(
                    w.tab.notebook.select(), text=os.path.basename(self.f.name)
                )
                return
        except AttributeError:
            return

    def save_file(self):
        try:
            text = w.tab.txt_collection[w.tab.notebook.index(w.tab.notebook.select())]
            self.f = open(self.f.name, "w")
            self.f.write(text.get("1.0", tk.END))
            self.f.close()
        except AttributeError:
            self.save_file_as()

    def new_file(self):
        txt, tb = w.tab.add_tab("untitled")
        txt.config(
            height=w.height,
            width=w.width,
        )
        txt.pack()
        tb.pack()

    def close_file(self):
        for current_tab in w.tab.notebook.winfo_children():
            if str(current_tab) == w.tab.notebook.select():
                w.tab.txt_collection.pop(w.tab.notebook.index(w.tab.notebook.select()))
                current_tab.destroy()
                return

    def exit(self):
        exit_win = tk.Toplevel(w.root)
        exit_win.title("")
        windowWidth = exit_win.winfo_reqwidth()
        windowHeight = exit_win.winfo_reqheight()
        positionRight = int(exit_win.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(exit_win.winfo_screenheight() / 2 - windowHeight / 2)
        exit_win.geometry(f"+{positionRight}+{positionDown}")
        exit_win.resizable(False, False)
        exit_label = tk.Label(exit_win, text="Are you sure you want to quit?")
        save_button = ttk.Button(
            exit_win, text="Save", command=lambda: self.save_file()
        )
        exit_button = ttk.Button(
            exit_win, text="Quit", command=lambda: w.root.destroy()
        )
        cancel_button = ttk.Button(
            exit_win, text="Cancel", command=lambda: exit_win.destroy()
        )
        exit_label.pack(side=tk.TOP)
        save_button.pack(side=tk.LEFT)
        exit_button.pack(side=tk.RIGHT)
        cancel_button.pack(side=tk.RIGHT)

    def binding_keys(self):
        w.root.bind_all("<Control-n>", lambda event: self.new_file())
        w.root.bind_all("<Control-N>", lambda event: self.new_file())
        w.root.bind_all("<Control-o>", lambda event: self.open_file())
        w.root.bind_all("<Control-O>", lambda event: self.open_file())
        w.root.bind_all("<Control-s>", lambda event: self.save_file())
        w.root.bind_all("<Control-S>", lambda event: self.save_file())
        w.root.bind_all("<Control-Shift-s>", lambda event: self.save_file_as())
        w.root.bind_all("<Control-Shift-S>", lambda event: self.save_file_as())
        w.root.bind_all("<Control-k>", lambda event: self.close_file())
        w.root.bind_all("<Control-K>", lambda event: self.close_file())

.pack()
        self.root.geometry("998x646")
        self.root.minsize(width=400, height=200)

    def footer_elements(self):
        while True:
            try:
                for i in w.tab.txt_collection:
                    word_count = len(i.get("1.0", tk.END).split())
                    character_count = len(i.get("1.0", tk.END)) - i.get(
                        "1.0", tk.END
                    ).count("\n")
                    mouse_pos = i.index(tk.INSERT)
                    line, column = (int(num) for num in mouse_pos.split("."))
                    self.footer.configure(
                        text=(
                            f"word count: {word_count}   "
                            f"character count: {character_count}     "
                            f"lines number: {line}    "
                            f"columns number: {column}"
                        ),
                        anchor=tk.E,
                    )
                self.root.update()
            except (tk.TclError, ValueError):
                break


w = Window()
import tkinter as tk
from tkinter import ttk
from textconfig import txtconfig
from customtext import CustomText


class Tab:
    def __init__(self, master):
        self.notebook = ttk.Notebook(master)
        self.txt_collection = []

    def add_tab(self, title):
        frame = tk.Frame(self.notebook, borderwidth=0, highlightthickness=0)
        self.vertical_scrollbar = tk.Scrollbar(frame)

        self.notebook.add(frame, text=title)
        self.txt = CustomText(
            frame,
            undo=True,
            insertbackground=txtconfig.insertbackground,
            background=txtconfig.background,
            foreground=txtconfig.foreground,
            selectbackground=txtconfig.selectbackground,
            yscrollcommand=self.vertical_scrollbar.set,
        )
        self.vertical_scrollbar.config(command=self.txt.yview, width=10)
        self.vertical_scrollbar.pack(fill="both", side=tk.RIGHT)
        self.txt.focus_set()
        self.txt.config(font=(txtconfig.font_family, txtconfig.font_size))
        self.txt_collection.append(self.txt)
        self.txt.bind("<Return>", self.indentation)
        return self.txt, self.notebook

    def indentation(self, event):
        self.txt = self.txt_collection[self.notebook.index(self.notebook.select())]
        try:
            pos = self.txt.index(tk.INSERT)
            line, column = (num for num in pos.split("."))
            func_len = (self.txt.get(f"{line}.0", f"{line}.{column}")).strip(" ")
            indentation_factor = abs(int(column) - len(func_len)) + 3
            auto_indent = abs(
                int(column)
                - len((self.txt.get(f"{line}.0", f"{line}.{column}").strip(" ")))
            )
            if ":" in self.txt.get("insert-1c"):
                self.txt.insert(tk.INSERT, "\n" + " " * indentation_factor)
                return "break"
            if "{" in self.txt.get("insert-1c"):
                self.txt.insert(tk.INSERT, "\n" + " " * indentation_factor)
                return "break"
            else:
                self.txt.insert(tk.INSERT, "\n" + " " * auto_indent)
                return "break"

        except AttributeError:
            return
import tkinter as tk
from tkinter import ttk
from window import w
from menubar import menubar


class Search:
    def __init__(self):
        self.search = tk.Menu(menubar.toolbar, tearoff=False)
        self.binding_keys()

    def search_menu(self):
        self.search.add_command(
            label="Find...", command=lambda: self.find_window(), accelerator="Ctrl+F"
        )
        self.search.add_command(
            label="Find and Replace",
            command=lambda: self.find_and_replace_window(),
            accelerator="Ctrl+H",
        )
        menubar.toolbar.add_cascade(label="Search", menu=self.search)
        w.root.config(menu=self.search)

    def find_window(self):
        self.find_win = tk.Toplevel(w.root)
        self.find_win.title("Find..")
        self.find_win.resizable(True, True)
        self.find_button = ttk.Button(
            self.find_win,
            text="           Find         ",
            command=lambda: self.find(),
        )
        self.find_all_button = ttk.Button(
            self.find_win,
            text="           Find all    ",
            command=lambda: self.find_all(),
        )
        self.entry = tk.Entry(self.find_win, width=75)
        self.find_win.geometry("400x100")

        self.entry.pack()
        self.find_all_button.pack(side=tk.LEFT, padx=(5, 0))
        self.find_button.pack(side=tk.LEFT, padx=(5, 0))

    def find(self):
        text = w.tab.txt_collection[(w.tab.notebook.index(w.tab.notebook.select()))]
        text.tag_remove("match", "1.0", tk.END)
        count = tk.IntVar()
        try:
            s = text.search(self.entry.get(), "1.0", count=count)
            text.tag_configure("match", background="orange")
            end = f"{s}+{count.get()}c"
            text.tag_add("match", s, end)
            text.bind(
                "<Button-1>", lambda event_data: text.tag_remove("match", "1.0", tk.END)
            )
        except IndexError:
            return
        if self.entry.get() in text.get("1.0", tk.END):
            return True
        else:
            return False

    def find_all(self):
        text = w.tab.txt_collection[w.tab.notebook.index(w.tab.notebook.select())]
        text.tag_remove("match", "1.0", tk.END)
        text.tag_configure("match", background="orange")
        count = tk.IntVar()
        addition_factor = 0
        index = "1.0"
        while True:
            try:
                s = text.search(
                    self.entry.get(), index, tk.END, count=count, regexp=True
                )
                end = f"{s}+{count.get()}c"
                text.tag_add("match", s, end)
                text.see(s)
                addition_factor += 1
                index = f"1.0+{addition_factor}c"
                text.bind(
                    "<Button-1>",
                    lambda event_data: w.txt.tag_remove("match", "1.0", tk.END),
                )
                w.root.update()
            except tk.TclError:
                break
        if self.entry.get() in text.get("1.0", tk.END):
            return True
        else:
            return False

    def find_and_replace_window(self):
        self.f_r_win = tk.Toplevel(w.root)
        self.f_r_win.title("Find and Replace")
        self.find_button2 = ttk.Button(
            self.f_r_win,
            text="          Find        ",
            command=lambda: self.find(),
        )
        self.find_all_button = ttk.Button(
            self.f_r_win,
            text="          Find all      ",
            command=lambda: self.find_all(),
        )
        self.entry = tk.Entry(self.f_r_win, width=75)
        self.f_r_win.geometry("600x240")
        self.f_r_win.resizable(False, False)
        self.replace_entry = tk.Entry(self.f_r_win, width=75)
        self.replace_button = ttk.Button(
            self.f_r_win,
            text="          Replace           ",
            command=lambda: self.replace(),
        )
        self.replace_all_button = ttk.Button(
            self.f_r_win,
            text="          Replace all       ",
            command=lambda: self.replace_all(),
        )
        self.entry.pack()
        self.replace_entry.pack()
        self.replace_all_button.pack(side=tk.LEFT, padx=(5, 0))
        self.replace_button.pack(side=tk.LEFT, padx=(5, 0))
        self.find_all_button.pack(side=tk.LEFT, padx=(5, 0))
        self.find_button2.pack(side=tk.LEFT, padx=(5, 0))

    def replace(self):
        text = w.tab.txt_collection[(w.tab.notebook.index(w.tab.notebook.select()))]
        try:
            if self.find():
                ranges = [str(i).strip("<>") for i in text.tag_ranges("match")]
                text.replace(ranges[0], ranges[-1], self.replace_entry.get())
        except (tk.TclError, IndexError):
            return

    def replace_all(self):
        text = w.tab.txt_collection[(w.tab.notebook.index(w.tab.notebook.select()))]
        try:
            if self.find_all():
                ranges = [str(i).strip("<>") for i in text.tag_ranges("match")]
                all_occurences = len(ranges) - 1
                while all_occurences > 0:
                    text.replace(
                        ranges[all_occurences - 1],
                        ranges[all_occurences],
                        self.replace_entry.get(),
                    )
                    all_occurences -= 2
        except tk.TclError:
            return

    def binding_keys(self):
        w.root.bind_all("<Control-f>", lambda event: self.find_window())
        w.root.bind_all("<Control-F>", lambda event: self.find_window())
        w.root.bind_all("<Control-h>", lambda event: self.find_and_replace_window())
        w.root.bind_all("<Control-H>", lambda event: self.find_and_replace_window())
import tkinter as tk
from tkinter import ttk
from window import w
from menubar import menubar


class Search:
    def __init__(self):
        self.search = tk.Menu(menubar.toolbar, tearoff=False)
        self.binding_keys()

    def search_menu(self):
        self.search.add_command(
            label="Find...", command=lambda: self.find_window(), accelerator="Ctrl+F"
        )
        self.search.add_command(
            label="Find and Replace",
            command=lambda: self.find_and_replace_window(),
            accelerator="Ctrl+H",
        )
        menubar.toolbar.add_cascade(label="Search", menu=self.search)
        w.root.config(menu=self.search)

    def find_window(self):
        self.find_win = tk.Toplevel(w.root)
        self.find_win.title("Find..")
        self.find_win.resizable(True, True)
        self.find_button = ttk.Button(
            self.find_win,
            text="           Find         ",
            command=lambda: self.find(),
        )
        self.find_all_button = ttk.Button(
            self.find_win,
            text="           Find all    ",
            command=lambda: self.find_all(),
        )
        self.entry = tk.Entry(self.find_win, width=75)
        self.find_win.geometry("400x100")

        self.entry.pack()
        self.find_all_button.pack(side=tk.LEFT, padx=(5, 0))
        self.find_button.pack(side=tk.LEFT, padx=(5, 0))

    def find(self):
        text = w.tab.txt_collection[(w.tab.notebook.index(w.tab.notebook.select()))]
        text.tag_remove("match", "1.0", tk.END)
        count = tk.IntVar()
        try:
            s = text.search(self.entry.get(), "1.0", count=count)
            text.tag_configure("match", background="orange")
            end = f"{s}+{count.get()}c"
            text.tag_add("match", s, end)
            text.bind(
                "<Button-1>", lambda event_data: text.tag_remove("match", "1.0", tk.END)
            )
        except IndexError:
            return
        if self.entry.get() in text.get("1.0", tk.END):
            return True
        else:
            return False

    def find_all(self):
        text = w.tab.txt_collection[w.tab.notebook.index(w.tab.notebook.select())]
        text.tag_remove("match", "1.0", tk.END)
        text.tag_configure("match", background="orange")
        count = tk.IntVar()
        addition_factor = 0
        index = "1.0"
        while True:
            try:
                s = text.search(
                    self.entry.get(), index, tk.END, count=count, regexp=True
                )
                end = f"{s}+{count.get()}c"
                text.tag_add("match", s, end)
                text.see(s)
                addition_factor += 1
                index = f"1.0+{addition_factor}c"
                text.bind(
                    "<Button-1>",
                    lambda event_data: w.txt.tag_remove("match", "1.0", tk.END),
                )
                w.root.update()
            except tk.TclError:
                break
        if self.entry.get() in text.get("1.0", tk.END):
            return True
        else:
            return False

    def find_and_replace_window(self):
        self.f_r_win = tk.Toplevel(w.root)
        self.f_r_win.title("Find and Replace")
        self.find_button2 = ttk.Button(
            self.f_r_win,
            text="          Find        ",
            command=lambda: self.find(),
        )
        self.find_all_button = ttk.Button(
            self.f_r_win,
            text="          Find all      ",
            command=lambda: self.find_all(),
        )
        self.entry = tk.Entry(self.f_r_win, width=75)
        self.f_r_win.geometry("600x240")
        self.f_r_win.resizable(False, False)
        self.replace_entry = tk.Entry(self.f_r_win, width=75)
        self.replace_button = ttk.Button(
            self.f_r_win,
            text="          Replace           ",
            command=lambda: self.replace(),
        )
        self.replace_all_button = ttk.Button(
            self.f_r_win,
            text="          Replace all       ",
            command=lambda: self.replace_all(),
        )
        self.entry.pack()
        self.replace_entry.pack()
        self.replace_all_button.pack(side=tk.LEFT, padx=(5, 0))
        self.replace_button.pack(side=tk.LEFT, padx=(5, 0))
        self.find_all_button.pack(side=tk.LEFT, padx=(5, 0))
        self.find_button2.pack(side=tk.LEFT, padx=(5, 0))

    def replace(self):
        text = w.tab.txt_collection[(w.tab.notebook.index(w.tab.notebook.select()))]
        try:
            if self.find():
                ranges = [str(i).strip("<>") for i in text.tag_ranges("match")]
                text.replace(ranges[0], ranges[-1], self.replace_entry.get())
        except (tk.TclError, IndexError):
            return

    def replace_all(self):
        text = w.tab.txt_collection[(w.tab.notebook.index(w.tab.notebook.select()))]
        try:
            if self.find_all():
                ranges = [str(i).strip("<>") for i in text.tag_ranges("match")]
                all_occurences = len(ranges) - 1
                while all_occurences > 0:
                    text.replace(
                        ranges[all_occurences - 1],
                        ranges[all_occurences],
                        self.replace_entry.get(),
                    )
                    all_occurences -= 2
        except tk.TclError:
            return

    def binding_keys(self):
        w.root.bind_all("<Control-f>", lambda event: self.find_window())
        w.root.bind_all("<Control-F>", lambda event: self.find_window())
        w.root.bind_all("<Control-h>", lambda event: self.find_and_replace_window())
        w.root.bind_all("<Control-H>", lambda event: self.find_and_replace_window())


class File:
    def __init__(self):
        self.filemenu = tk.Menu(menubar.toolbar, tearoff=False)
        self.binding_keys()

    def file_menu(self):
        self.filemenu.add_command(
            label="New file", command=lambda: self.new_file(), accelerator="Ctrl+N"
        )
        self.filemenu.add_command(
            label="Open File", command=lambda: self.open_file(), accelerator="Ctrl+O"
        )
        self.filemenu.add_command(
            label="Save", command=lambda: self.save_file(), accelerator="Ctrl+S"
        )
        self.filemenu.add_command(
            label="Save as",
            command=lambda: self.save_file_as(),
            accelerator="Ctrl+Shift+S",
        )
        self.filemenu.add_command(
            label="Close file", command=lambda: self.close_file(), accelerator="Ctrl+K"
        )
        self.filemenu.add_command(label="Exit", command=lambda: self.exit())
        menubar.toolbar.add_cascade(label="File", menu=self.filemenu)

    def open_file(self):
        try:
            self.f = filedialog.askopenfile(
                mode="r",
                initialdir="/",
                title="Select A File",
                filetypes=(("files", "*.txt"), ("all files", "*.*")),
            )
            filename = os.path.basename(self.f.name)
            txt, tb = w.tab.add_tab(f"{filename}")
            txt.config(
                height=w.height,
                width=w.width,
            )
            txt.insert("1.0", self.f.read())
            txt.pack()
            tb.pack()

        except AttributeError:
            return

    def save_file_as(self):
        try:
            for i in w.tab.txt_collection:
                self.f = filedialog.asksaveasfile(
                    mode="w",
                    defaultextension=".txt",
                    initialdir="/",
                    title="Select A File",
                    filetypes=(
                        ("text", "*.txt"),
                        ("python", "*.py"),
                        ("all files", "*.*"),
                    ),
                )
                self.f.write(i.get("1.0", tk.END))
                self.f.close()
                w.tab.notebook.tab(
                    w.tab.notebook.select(), text=os.path.basename(self.f.name)
                )
                return
        except AttributeError:
            return

    def save_file(self):
        try:
            text = w.tab.txt_collection[w.tab.notebook.index(w.tab.notebook.select())]
            self.f = open(self.f.name, "w")
            self.f.write(text.get("1.0", tk.END))
            self.f.close()
        except AttributeError:
            self.save_file_as()

    def new_file(self):
        txt, tb = w.tab.add_tab("untitled")
        txt.config(
            height=w.height,
            width=w.width,
        )
        txt.pack()
        tb.pack()

    def close_file(self):
        for current_tab in w.tab.notebook.winfo_children():
            if str(current_tab) == w.tab.notebook.select():
                w.tab.txt_collection.pop(w.tab.notebook.index(w.tab.notebook.select()))
                current_tab.destroy()
                return

    def exit(self):
        exit_win = tk.Toplevel(w.root)
        exit_win.title("")
        windowWidth = exit_win.winfo_reqwidth()
        windowHeight = exit_win.winfo_reqheight()
        positionRight = int(exit_win.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(exit_win.winfo_screenheight() / 2 - windowHeight / 2)
        exit_win.geometry(f"+{positionRight}+{positionDown}")
        exit_win.resizable(False, False)
        exit_label = tk.Label(exit_win, text="Are you sure you want to quit?")
        save_button = ttk.Button(
            exit_win, text="Save", command=lambda: self.save_file()
        )
        exit_button = ttk.Button(
            exit_win, text="Quit", command=lambda: w.root.destroy()
        )
        cancel_button = ttk.Button(
            exit_win, text="Cancel", command=lambda: exit_win.destroy()
        )
        exit_label.pack(side=tk.TOP)
        save_button.pack(side=tk.LEFT)
        exit_button.pack(side=tk.RIGHT)
        cancel_button.pack(side=tk.RIGHT)

    def binding_keys(self):
        w.root.bind_all("<Control-n>", lambda event: self.new_file())
        w.root.bind_all("<Control-N>", lambda event: self.new_file())
        w.root.bind_all("<Control-o>", lambda event: self.open_file())
        w.root.bind_all("<Control-O>", lambda event: self.open_file())
        w.root.bind_all("<Control-s>", lambda event: self.save_file())
        w.root.bind_all("<Control-S>", lambda event: self.save_file())
        w.root.bind_all("<Control-Shift-s>", lambda event: self.save_file_as())
        w.root.bind_all("<Control-Shift-S>", lambda event: self.save_file_as())
        w.root.bind_all("<Control-k>", lambda event: self.close_file())
        w.root.bind_all("<Control-K>", lambda event: self.close_file())




class Edit:
    def __init__(self):
        self.edit = tk.Menu(menubar.toolbar, tearoff=False)

    def edit_menu(self):
        self.edit.add_command(
            label="Copy", command=lambda: self.copy(), accelerator="Ctrl+C"
        )
        self.edit.add_command(
            label="Paste", command=lambda: self.paste(), accelerator="Ctrl+V"
        )
        self.edit.add_command(
            label="Cut", command=lambda: self.cut(), accelerator="Ctrl+X"
        )
        self.edit.add_command(
            label="Redo", command=lambda: self.redo(), accelerator="Ctrl+Y"
        )
        self.edit.add_command(
            label="Undo", command=lambda: self.undo(), accelerator="Ctrl+Z"
        )
        self.edit.add_command(
            label="Select all", command=lambda: self.select_all(), accelerator="Ctrl+A"
        )
        self.edit.add_command(
            label="Delete", command=lambda: keyboard.press("DEL"), accelerator="DEL"
        )
        menubar.toolbar.add_cascade(label="Edit", menu=self.edit)
        w.root.config(menu=self.edit)

    def copy(self):
        for i in w.tab.txt_collection:
            i.event_generate("<<Copy>>")
            w.root.update()

    def paste(self):
        for i in w.tab.txt_collection:
            i.event_generate("<<Paste>>")
            w.root.update()

    def cut(self):
        for i in w.tab.txt_collection:
            i.event_generate("<<Cut>>")
            w.root.update()

    def redo(self):
        for i in w.tab.txt_collection:
            i.event_generate("<<Redo>>")
            w.root.update()

    def undo(self):
        for i in w.tab.txt_collection:
            i.event_generate("<<Undo>>")
            w.root.update()

    def select_all(self):
        w.root.focus_get().tag_add("sel", "1.0", tk.END)
import tkinter as tk


class CustomText(tk.Text):
    def __init__(self, *args, **kwargs):
        tk.Text.__init__(self, *args, **kwargs)

        # create a proxy for the underlying widget
        self._orig = self._w + "_orig"
        self.tk.call("rename", self._w, self._orig)
        self.tk.createcommand(self._w, self._proxy)

    def _proxy(self, *args):
        try:
            # let the actual widget perform the requested action
            cmd = (self._orig,) + args
            result = self.tk.call(cmd)
            # generate an event if something was added or deleted,
            # or the cursor position changed
            if (
                args[0] in ("insert", "replace", "delete")
                or args[0:3] == ("mark", "set", "insert")
                or args[0:2] == ("xview", "moveto")
                or args[0:2] == ("xview", "scroll")
                or args[0:2] == ("yview", "moveto")
                or args[0:2] == ("yview", "scroll")
            ):  # return indentation to original amount
                self.event_generate("<<Change>>", when="tail")

            # return what the actual widget returned
            return result
        except tk.TclError:
            return



class View:
    def __init__(self):
        self.view = tk.Menu(menubar.toolbar, tearoff=False)
        self.font_size = tk.Menu(self.view, tearoff=False)
        self.font_family = tk.Menu(self.view, tearoff=False)

    def view_menu(self):
        self.set_size_command()
        self.set_family_command()
        self.view.add_cascade(label="Font size", menu=self.font_size)
        self.view.add_cascade(label="Font Family", menu=self.font_family)
        menubar.toolbar.add_cascade(label="View", menu=self.view)
        w.root.config(menu=self.view)

    def size_and_family(self, fm, px):
        txtconfig.font_family = fm
        txtconfig.font_size = px
        fonts = font.Font(family=txtconfig.font_family, size=txtconfig.font_size)
        for i in w.tab.txt_collection:
            i.config(font=(fonts))
            w.root.update()

    def set_size_command(self):
        self.font_size.add_radiobutton(
            label="12px(Default)",
            command=lambda: self.size_and_family(txtconfig.font_family, 12),
        )
        self.font_size.add_radiobutton(
            label="15px",
            command=lambda: self.size_and_family(txtconfig.font_family, 15),
        )
        self.font_size.add_radiobutton(
            label="20px",
            command=lambda: self.size_and_family(txtconfig.font_family, 20),
        )
        self.font_size.add_radiobutton(
            label="35px", command=lambda: self.size_and_family(txtconfig, 35)
        )
        self.font_size.add_radiobutton(
            label="45px",
            command=lambda: self.size_and_family(txtconfig.font_family, 45),
        )

    def set_family_command(self):
        self.font_family.add_radiobutton(
            label="Consolas",
            command=lambda: self.size_and_family("Consolas", txtconfig.font_size),
        )
        self.font_family.add_radiobutton(
            label="Arial",
            command=lambda: self.size_and_family("Arial", txtconfig.font_size),
        )
        self.font_family.add_radiobutton(
            label="Modern",
            command=lambda: self.size_and_family("Modern", txtconfig.font_size),
        )
        self.font_family.add_radiobutton(
            label="Roman",
            command=lambda: self.size_and_family("Roman", txtconfig.font_size),
        )
        self.font_family.add_radiobutton(
            label="Courier",
            command=lambda: self.size_and_family("Courier", txtconfig.font_size),
        )
        self.font_family.add_radiobutton(
            label="Terminal",
            command=lambda: self.size_and_family("Terminal", txtconfig.font_size),
        )
        self.font_family.add_radiobutton(
            label="Palatino Linotype",
            command=lambda: self.size_and_family(
                "Palatino Linotype", txtconfig.font_size
            ),
        )
"""Made  a class for text configuration, so it is possible and easier to
control text configurations around all widgets at once,
also can be used to edit some other widgets background"""


class TextConfig:
    def __init__(self):
        self.insertbackground = "white"
        self.background = "#282828"
        self.foreground = "white"
        self.selectbackground = "blue"
        self.font_family = "Consolas"
        self.font_size = 12


txtconfig = TextConfig()

import webbrowser



import webbrowser


import subprocess


class Run:
    def __init__(self):
        self.file = File()
        self.file.file_menu()
        self.run = tk.Menu(menubar.toolbar, tearoff=False)
        self.run.add_command(label="Run", command=lambda: self.run_it())
        menubar.toolbar.add_cascade(label="Run", menu=self.run)
        self.binding_keys()

    def run_it(self):
        self.file.save_file()
        try:
            if self.file.f.name.endswith(".py"):
                subprocess.Popen(["python", self.file.f.name], shell=True)
            elif self.file.f.name.endswith(".js"):
                subprocess.Popen(["node", self.file.f.name], shell=True)
            elif self.file.f.name.endswith(".exs"):
                subprocess.Popen(["elixir", self.file.f.name], shell=True)
            else:
                toplevel = tk.Toplevel(w.root)
                toplevel.title("Warning")
                txt = tk.Text(toplevel)
                txt.insert("1.0", "File cannot run.")
                toplevel.geometry("215x50")
                toplevel.resizable(False, False)
                txt.config(state=tk.DISABLED)
                txt.pack()
        except AttributeError:
            return

    def binding_keys(self):
        w.root.bind_all("<F5>", lambda event: self.run_it())


"""bind a key for run,
and it should check file extension,
to run it in the suitable programming language."""
# Running a program need a file, so the run and compile file has File object




class Tools:
    def __init__(self):
        self.tool_menu = tk.Menu(menubar.toolbar, tearoff=False)

    def tools_menu(self):
        self.tool_menu.add_command(
            accelerator="---------"
        )
        self.tool_menu.add_command(
            accelerator="-----"
        )
        menubar.toolbar.add_cascade(label="Tools", menu=self.tool_menu)
        w.root.config(menu=self.tool_menu)

    def google_search_win(self):
        toplevel = tk.Toplevel(w.root)
        toplevel.title("")
        toplevel.geometry("200x100")
        toplevel.resizable(False, False)
        self.entry = tk.Entry(toplevel, width=100)
        search_button = ttk.Button(
            toplevel, text="Google search"
        )
        search_button.pack(side=tk.BOTTOM, padx=(5, 0))
        self.entry.pack()

import tkinter as tk

# Sample boilerplate code snippets
BOILERPLATE_CODE = {
    "if statement": "if condition: \n "+"    " ,
    "for loop": "for item in items: \n" +"    ",
    "function": "def my_function(arg1, arg2): \n" + "    ",
    "imports": "import tkinter as tk \n" + "    ",
}


def insert_snippet():
    selected_snippet = combo_var.get()
    snippet_code = BOILERPLATE_CODE.get(selected_snippet, "")
    text_widget.insert("insert", snippet_code)
import docx
import tkinter as tk
from tkinter import filedialog, Text, scrolledtext

def read_word_file(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def load_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
    if file_path:
        content = read_word_file(file_path)
        text_box.delete(1.0, tk.END)
        text_box.insert(tk.END, content)

# Create the main window
app = tk.Tk()
app.title("Word Document Reader")

# Add a button to load the Word document
load_button = tk.Button(app, text="Load Word Document", command=load_file)
load_button.pack(pady=20)

# Add a scrolled text widget to display the content
text_box = scrolledtext.ScrolledText(app, wrap=tk.WORD)
text_box.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

app.mainloop()


# Create the main application
app = tk.Tk()
app.title("Auto Coder")

# Create a combo box to select snippets
combo_var = tk.StringVar(app)
combo_box = tk.OptionMenu(app, combo_var, *BOILERPLATE_CODE.keys())
combo_box.pack(pady=10)

# Create the text widget for code editing
text_widget = tk.Text(app, wrap="word", width=50, height=20)
text_widget.pack()

# Create the insert button
insert_button = tk.Button(app, text="Insert Snippet", command=insert_snippet)
insert_button.pack(pady=10)

app.mainloop()


import docx
import tkinter as tk
from tkinter import filedialog, Text, scrolledtext

def read_word_file(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def load_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
    if file_path:
        content = read_word_import docx
import tkinter as tk
from tkinter import filedialog, Text, scrolledtext

def read_word_file(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def load_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
    if file_path:
        content = read_word_file(file_path)
        text_box.delete(1.0, tk.END)
        text_box.insert(tk.END, content)

# Create the main window
app = tk.Tk()
app.title("Word Document Reader")

import tkinter as tk
from tkinter import ttk, filedialog

class UnicodeListGeneratorApp:
    def __init__(self, master):
        self.master = master
        self.master.title("Unicode List Generator")
        
        # Create and place labels, entries, and button
        start_label = ttk.Label(self.master, text="Start (hex):")
        start_label.grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)

        end_label = ttk.Label(self.master, text="End (hex):")
        end_label.grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)

        self.start_entry = ttk.Entry(self.master, width=10)
        self.start_entry.grid(row=0, column=1, padx=10, pady=5)

        self.end_entry = ttk.Entry(self.master, width=10)
        self.end_entry.grid(row=1, column=1, padx=10, pady=5)

        generate_button = ttk.Button(self.master, text="Generate", command=self.generate_unicode)
        generate_button.grid(row=2, column=0, columnspan=2, pady=10)

        save_button = ttk.Button(self.master, text="Save to File", command=self.save_to_file)
        save_button.grid(row=2, column=4)

        clear_button = ttk.Button(self.master, text="Clear list", command=self.clear)
        clear_button.grid(row=2, column=6, pady=10)

        # Create and place listbox for output
        self.output_listbox = tk.Listbox(self.master, width=20, height=15)
        self.output_listbox.grid(row=3, column=0, columnspan=2, padx=10, pady=5)

    def generate_unicode(self):
        start_hex = self.start_entry.get()
        end_hex = self.end_entry.get()

        # Convert hex strings to integers
        start_int = int(start_hex, 16)
        end_int = int(end_hex, 16)

        # Clear the listbox
        self.output_listbox.delete(0, tk.END)

        # Add characters to the listbox
        for code_point in range(start_int, end_int + 1):
            char = chr(code_point)
            hex_value = f"{code_point:04X}"
            self.output_listbox.insert(tk.END, f"U+{hex_value}: {char}")

    def save_to_file(self):
        # Choose where to save the file
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                  filetypes=[("Text files", "*.txt"),
                                                             ("All files", "*.*")])
        if not file_path:  # if the user cancels the save
            return

        with open(file_path, 'w', encoding='utf-8') as file:
            for item in self.output_listbox.get(0, tk.END):
                file.write(item + ',' + '\n')

    def clear(self):
        self.output_listbox.delete(0, tk.END)

if __name__ == '__main__':
    root = tk.Tk()
    app = UnicodeListGeneratorApp(root)
    root.mainloop()
# Add a button to load the Word document
load_button = tk.Button(app, text="Load Word Document", command=load_file)
load_button.pack(pady=20)

# Add a scrolled text widget to display the content
text_box = scrolledtext.ScrolledText(app, wrap=tk.WORD)
text_box.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

app.mainloop()
file(file_path)
        text_box.delete(1.0, tk.END)
        text_box.insert(tk.END, content)

# Create the main window
app = tk.Tk()
app.title("Word Document Reader")

# Add a button to load the Word document
load_button = tk.Button(app, text="Load Word Document", command=load_file)
load_button.pack(pady=20)

# Add a scrolled text widget to display the content
text_box = scrolledtext.ScrolledText(app, wrap=tk.WORD)
text_box.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

app.mainloop()
import json

def process_text_file(filename):
    with open(filename, 'r') as f:
        lines = f.readlines()
    
    # Parse the key-value pairs and add double quotes around the values
    data = {}
    for line in lines:
        key, value = line.strip().split(":")
        data[key.strip()] = f'"{value.strip()}"'  # Double-quote the value
        
    return data

def save_as_json(data, filename):
    with open(filename, 'w') as f:
        json.dump(data, f, indent=4)

def open_file_dialog():
    filename = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if filename:
        try:
            data = process_text_file(filename)
            json_filename = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON files", "*.json")])
            if json_filename:
                save_as_json(data, json_filename)
                messagebox.showinfo("Success", "Text file processed and saved as JSON successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

# Set up the GUI
root = tk.Tk()
root.title("Convert Text to JSON")
root.geometry("350x120")

label = tk.Label(root, text="Select a text file to convert to JSON")
label.pack(pady=10)

button = tk.Button(root, text="Select File", command=open_file_dialog)
button.pack(pady=10)

root.mainloop()

